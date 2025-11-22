#!/usr/bin/env python3
"""
Convert all markdown documentation to Word (.docx) format.

Requires: pip install python-docx markdown
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE


def md_to_docx(md_file, output_dir):
    """Convert a markdown file to Word document."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set document properties
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Process markdown content
    lines = content.split('\n')
    current_list = None
    in_code_block = False
    code_content = []

    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, 0)
            elif level == 2:
                p = doc.add_heading(text, 1)
            elif level == 3:
                p = doc.add_heading(text, 2)
            elif level == 4:
                p = doc.add_heading(text, 3)
            else:
                p = doc.add_heading(text, 4)
            continue

        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            continue

        # Handle checkboxes
        if line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
            checked = '[x]' in line
            text = line.strip()[6:]
            marker = '☑' if checked else '☐'
            p = doc.add_paragraph(f'{marker} {text}')
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Handle tables (basic)
        if '|' in line and not line.strip().startswith('|--'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                # Simple table representation
                p = doc.add_paragraph('  |  '.join(cells))
                p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(0)
            for run in p.runs:
                run.italic = True
            continue

        # Handle bold and italic inline
        text = line
        if text.strip():
            p = doc.add_paragraph()
            # Process inline formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)
        else:
            # Empty line - add paragraph break
            doc.add_paragraph()

    # Save document
    base_name = Path(md_file).stem
    output_file = os.path.join(output_dir, f'{base_name}.docx')
    doc.save(output_file)
    return output_file


def main():
    """Convert all markdown files in docs directory to Word format."""
    docs_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_dir = os.path.join(docs_dir, 'word_documents')

    # Create output directory
    os.makedirs(output_dir, exist_ok=True)

    # Find all markdown files
    md_files = []
    for f in os.listdir(docs_dir):
        if f.endswith('.md'):
            md_files.append(os.path.join(docs_dir, f))

    print(f"Found {len(md_files)} markdown files to convert")

    # Convert each file
    for md_file in md_files:
        try:
            output_file = md_to_docx(md_file, output_dir)
            print(f"✓ Converted: {os.path.basename(md_file)} -> {os.path.basename(output_file)}")
        except Exception as e:
            print(f"✗ Failed: {os.path.basename(md_file)} - {str(e)}")

    print(f"\nAll Word documents saved to: {output_dir}")


if __name__ == "__main__":
    main()
